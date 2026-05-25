"""
文档上传和文本处理管线（ai/rag/pipeline.py）

说明：实现 RAG 文档处理的完整管线，包括文档加载、文本清洗、
     递归字符分块、向量嵌入生成和数据库存储。

处理流程：
    上传文件 → 提取文本 → 清洗文本 → 递归分块 → 生成嵌入 → 存储到 pgvector

支持的文件类型：
    - PDF: 使用 PyPDF2 提取文本
    - DOCX: 使用 python-docx 提取文本
    - MD/TXT: 直接读取文本内容

分块策略：
    使用递归字符分割器，按照分隔符优先级递归切分文本：
    1. 双换行符（段落分隔）
    2. 单换行符（行分隔）
    3. 句号/问号/感叹号（句子分隔）
    4. 逗号/分号（子句分隔）
    5. 空格（单词分隔）
    6. 单字符（最后手段）

用法：
    from app.ai.rag.pipeline import ingest_document

    result = await ingest_document(
        file_path="/tmp/upload/doc.pdf",
        filename="公司制度.pdf",
        file_type="pdf",
        db=session,
    )
"""

import logging
import re
from datetime import datetime
from pathlib import Path

import httpx
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai.rag.models import RagChunk, RagDocument
from app.core.config import settings

logger = logging.getLogger(__name__)


# ============================================================
# 文档处理管线主入口
# ============================================================


async def ingest_document(
    file_path: str,
    filename: str,
    file_type: str,
    file_size: int,
    db: AsyncSession,
) -> dict:
    """
    文档处理管线主入口

    说明：完整的文档处理流程，从文件读取到向量存储。
         处理过程中更新文档状态，便于前端展示进度。

    参数：
        file_path: 临时文件路径
        filename: 原始文件名
        file_type: 文件类型（pdf/docx/md/txt）
        file_size: 文件大小（字节）
        db: PostgreSQL 异步数据库会话

    返回：
        {"doc_id": int, "filename": str, "chunks": int, "status": str}

    异常：
        处理失败时更新文档状态为 "failed"，不抛出异常到上层
    """
    # 1. 创建文档记录，状态为 processing
    document = RagDocument(
        filename=filename,
        file_type=file_type,
        file_size=file_size,
        status="processing",
    )
    db.add(document)
    await db.flush()  # 获取 doc_id，但不提交事务

    try:
        # 2. 根据文件类型提取文本
        raw_text = await extract_text(file_path, file_type)
        if not raw_text or len(raw_text.strip()) < 10:
            raise ValueError(f"文档内容为空或过短: {filename}")

        # 3. 清洗文本
        cleaned_text = clean_text(raw_text)
        if not cleaned_text:
            raise ValueError(f"清洗后文本为空: {filename}")

        # 4. 递归字符分块
        chunks = split_text(
            text=cleaned_text,
            chunk_size=settings.RAG_CHUNK_SIZE,
            chunk_overlap=settings.RAG_CHUNK_OVERLAP,
        )
        if not chunks:
            raise ValueError(f"分块结果为空: {filename}")

        # 5. 生成向量嵌入
        embeddings = await generate_embeddings(chunks)

        # 6. 存储分块到数据库
        for idx, (chunk_text, embedding) in enumerate(zip(chunks, embeddings)):
            chunk = RagChunk(
                doc_id=document.doc_id,
                chunk_index=idx,
                content=chunk_text,
                embedding=embedding,
                token_count=estimate_token_count(chunk_text),
            )
            db.add(chunk)

        # 7. 更新文档状态为 ready
        document.status = "ready"
        document.chunk_count = len(chunks)
        document.processed_time = datetime.now()

        await db.flush()

        logger.info(f"文档处理完成: {filename}, 分块数: {len(chunks)}")
        return {
            "doc_id": document.doc_id,
            "filename": filename,
            "chunks": len(chunks),
            "status": "ready",
        }

    except Exception as e:
        # 处理失败，更新状态
        document.status = "failed"
        document.processed_time = datetime.now()
        await db.flush()

        logger.error(f"文档处理失败: {filename}, 错误: {e}")
        return {
            "doc_id": document.doc_id,
            "filename": filename,
            "chunks": 0,
            "status": "failed",
            "error": str(e),
        }


# ============================================================
# 文本提取
# ============================================================


async def extract_text(file_path: str, file_type: str) -> str:
    """
    根据文件类型提取文本内容

    参数：
        file_path: 文件路径
        file_type: 文件类型（pdf/docx/md/txt）

    返回：
        提取的原始文本内容
    """
    path = Path(file_path)

    if file_type == "pdf":
        return _extract_pdf(path)
    elif file_type == "docx":
        return _extract_docx(path)
    elif file_type in ("md", "txt"):
        return _extract_plain_text(path)
    else:
        # 尝试作为纯文本读取
        return _extract_plain_text(path)


def _extract_pdf(path: Path) -> str:
    """
    从 PDF 文件提取文本

    说明：使用 PyPDF2 逐页提取文本。如果 PyPDF2 不可用，
         回退到以纯文本方式读取（可能乱码）。
    """
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        return "\n\n".join(pages_text)
    except ImportError:
        logger.warning("PyPDF2 未安装，尝试以纯文本方式读取 PDF")
        return _extract_plain_text(path)
    except Exception as e:
        logger.error(f"PDF 提取失败: {e}")
        raise ValueError(f"PDF 文本提取失败: {e}")


def _extract_docx(path: Path) -> str:
    """
    从 DOCX 文件提取文本

    说明：使用 python-docx 逐段落提取文本。如果 python-docx 不可用，
         回退到以纯文本方式读取。
    """
    try:
        from docx import Document

        doc = Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        return "\n\n".join(paragraphs)
    except ImportError:
        logger.warning("python-docx 未安装，尝试以纯文本方式读取 DOCX")
        return _extract_plain_text(path)
    except Exception as e:
        logger.error(f"DOCX 提取失败: {e}")
        raise ValueError(f"DOCX 文本提取失败: {e}")


def _extract_plain_text(path: Path) -> str:
    """从纯文本文件读取内容（支持 UTF-8 和 GBK 编码）"""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    raise ValueError(f"无法解码文件: {path.name}")


# ============================================================
# 文本清洗
# ============================================================


def clean_text(raw_text: str) -> str:
    """
    清洗原始文本

    处理步骤：
        1. 移除页眉页脚标记（常见的页码格式）
        2. 规范化空白字符（多个空格合并为一个）
        3. 移除过短的行（少于 20 个字符，通常是页码或标注）
        4. 移除连续的空行（保留最多一个空行）

    参数：
        raw_text: 原始提取的文本

    返回：
        清洗后的文本
    """
    if not raw_text:
        return ""

    # 移除常见的页眉页脚模式
    # 例如: "第 1 页 共 10 页", "Page 1 of 10", "- 1 -"
    text = re.sub(r"第\s*\d+\s*页\s*共\s*\d+\s*页", "", raw_text)
    text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.IGNORECASE)
    text = re.sub(r"[-—]\s*\d+\s*[-—]", "", text)

    # 规范化空白字符
    # 将制表符和多个空格替换为单个空格
    text = re.sub(r"[ \t]+", " ", text)

    # 按行处理，移除过短的行
    lines = text.split("\n")
    filtered_lines = []
    for line in lines:
        stripped = line.strip()
        # 保留空行（用于段落分隔）和长度 >= 20 的行
        if stripped == "" or len(stripped) >= 20:
            filtered_lines.append(stripped)

    # 合并连续空行为单个空行
    result_lines = []
    prev_empty = False
    for line in filtered_lines:
        if line == "":
            if not prev_empty:
                result_lines.append("")
            prev_empty = True
        else:
            result_lines.append(line)
            prev_empty = False

    return "\n".join(result_lines).strip()


# ============================================================
# 递归字符分块器
# ============================================================

# 分隔符优先级列表（从粗粒度到细粒度）
_SEPARATORS = [
    "\n\n",  # 段落分隔
    "\n",  # 行分隔
    "。",  # 中文句号
    "！",  # 中文感叹号
    "？",  # 中文问号
    ".",  # 英文句号
    "!",  # 英文感叹号
    "?",  # 英文问号
    "；",  # 中文分号
    ";",  # 英文分号
    "，",  # 中文逗号
    ",",  # 英文逗号
    " ",  # 空格
    "",  # 单字符（最后手段）
]


def split_text(
    text: str,
    chunk_size: int = 512,
    chunk_overlap: int = 50,
) -> list[str]:
    """
    递归字符分块器

    说明：按照分隔符优先级递归切分文本，确保每个分块不超过 chunk_size。
         相邻分块之间有 chunk_overlap 个字符的重叠，保证上下文连续性。

    算法：
        1. 从最粗粒度的分隔符开始尝试切分
        2. 如果切分后的片段仍然超过 chunk_size，使用更细粒度的分隔符递归切分
        3. 最终确保每个分块长度在 [1, chunk_size] 范围内
        4. 添加重叠部分保证语义连续

    参数：
        text: 待分块的文本
        chunk_size: 每个分块的最大字符数
        chunk_overlap: 相邻分块的重叠字符数

    返回：
        分块后的文本列表
    """
    if not text:
        return []

    # 如果文本长度不超过 chunk_size，直接返回
    if len(text) <= chunk_size:
        return [text]

    # 递归分割
    chunks = _recursive_split(text, _SEPARATORS, chunk_size)

    # 添加重叠
    if chunk_overlap > 0 and len(chunks) > 1:
        chunks = _add_overlap(chunks, chunk_overlap)

    # 过滤空分块
    return [chunk.strip() for chunk in chunks if chunk.strip()]


def _recursive_split(
    text: str,
    separators: list[str],
    chunk_size: int,
) -> list[str]:
    """
    递归分割文本

    说明：尝试使用当前分隔符切分文本，如果切分后的片段仍然过长，
         使用下一级分隔符继续递归切分。
    """
    if not text or len(text) <= chunk_size:
        return [text] if text else []

    if not separators:
        # 没有更多分隔符，强制按 chunk_size 切分
        return _force_split(text, chunk_size)

    separator = separators[0]
    remaining_separators = separators[1:]

    # 使用当前分隔符切分
    if separator == "":
        # 空字符串分隔符 = 逐字符切分
        return _force_split(text, chunk_size)

    parts = text.split(separator)

    # 合并小片段，确保不超过 chunk_size
    chunks = []
    current_chunk = ""

    for part in parts:
        # 计算加入当前片段后的长度
        candidate = current_chunk + separator + part if current_chunk else part

        if len(candidate) <= chunk_size:
            current_chunk = candidate
        else:
            # 当前块已满，保存并开始新块
            if current_chunk:
                chunks.append(current_chunk)

            # 如果单个 part 超过 chunk_size，递归使用更细的分隔符
            if len(part) > chunk_size:
                sub_chunks = _recursive_split(part, remaining_separators, chunk_size)
                chunks.extend(sub_chunks)
                current_chunk = ""
            else:
                current_chunk = part

    # 保存最后一个块
    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def _force_split(text: str, chunk_size: int) -> list[str]:
    """强制按固定长度切分文本（最后手段）"""
    chunks = []
    for i in range(0, len(text), chunk_size):
        chunk = text[i : i + chunk_size]
        if chunk:
            chunks.append(chunk)
    return chunks


def _add_overlap(chunks: list[str], overlap: int) -> list[str]:
    """
    为相邻分块添加重叠部分

    说明：每个分块的开头包含前一个分块末尾的 overlap 个字符，
         确保语义上下文的连续性。
    """
    if not chunks or overlap <= 0:
        return chunks

    result = [chunks[0]]
    for i in range(1, len(chunks)):
        prev_chunk = chunks[i - 1]
        # 取前一个分块末尾的 overlap 个字符作为当前分块的前缀
        overlap_text = prev_chunk[-overlap:] if len(prev_chunk) >= overlap else prev_chunk
        result.append(overlap_text + chunks[i])

    return result


# ============================================================
# 向量嵌入生成
# ============================================================


async def generate_embeddings(texts: list[str]) -> list[list[float]]:
    """
    生成文本向量嵌入

    说明：调用嵌入模型 API 生成文本的向量表示。
         如果 API 未配置或调用失败，回退到零向量占位。

    嵌入 API 兼容 OpenAI 格式：
        POST /embeddings
        {"model": "text-embedding-v3", "input": ["text1", "text2"]}

    参数：
        texts: 待嵌入的文本列表

    返回：
        向量列表，每个向量为 EMBEDDING_DIMENSIONS 维浮点数组
    """
    dimensions = settings.EMBEDDING_DIMENSIONS

    # 检查 API 是否已配置
    api_key = settings.embedding_api_key
    if not api_key or api_key == "your_dashscope_api_key_here":
        logger.warning("嵌入 API 未配置，使用零向量占位")
        return [[0.0] * dimensions for _ in texts]

    try:
        # 分批处理（每批最多 20 条，避免超出 API 限制）
        all_embeddings = []
        batch_size = 20

        for i in range(0, len(texts), batch_size):
            batch = texts[i : i + batch_size]
            batch_embeddings = await _call_embedding_api(batch)
            all_embeddings.extend(batch_embeddings)

        return all_embeddings

    except Exception as e:
        logger.error(f"嵌入 API 调用失败，使用零向量占位: {e}")
        return [[0.0] * dimensions for _ in texts]


async def _call_embedding_api(texts: list[str]) -> list[list[float]]:
    """
    调用嵌入模型 API（兼容 OpenAI 格式）

    说明：使用 httpx 异步 HTTP 客户端调用嵌入 API。
         API 格式兼容 OpenAI /v1/embeddings 接口。
    """
    url = f"{settings.EMBEDDING_BASE_URL}/embeddings"
    headers = {
        "Authorization": f"Bearer {settings.embedding_api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": settings.EMBEDDING_MODEL,
        "input": texts,
        "dimensions": settings.EMBEDDING_DIMENSIONS,
    }

    async with httpx.AsyncClient(timeout=60.0) as client:
        response = await client.post(url, json=payload, headers=headers)
        response.raise_for_status()

        data = response.json()
        # OpenAI 格式: {"data": [{"embedding": [...], "index": 0}, ...]}
        embeddings = [item["embedding"] for item in data["data"]]
        return embeddings


# ============================================================
# 辅助函数
# ============================================================


def estimate_token_count(text: str) -> int:
    """
    估算文本的 Token 数量

    说明：简单估算方法：
         - 中文：约 1 个字符 = 1-2 个 Token
         - 英文：约 4 个字符 = 1 个 Token
         这里使用混合估算：总字符数 * 0.6

    参数：
        text: 待估算的文本

    返回：
        估算的 Token 数量
    """
    if not text:
        return 0

    # 统计中文字符数和非中文字符数
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", text))
    other_chars = len(text) - chinese_chars

    # 中文约 1.5 token/字符，英文约 0.25 token/字符
    return int(chinese_chars * 1.5 + other_chars * 0.25)
